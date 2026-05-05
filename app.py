import streamlit as st
import uuid
import pandas as pd
from database import init_db, save_final_report, SessionLocal, AgentOutput, FinalReport, User, get_user_by_username, create_user, update_report_rating
from samples import INCIDENTS
from agents import run_rca_workflow
from docx import Document
from io import BytesIO
import re
import logging
import traceback
import sys
import threading
import time
import streamlit_authenticator as stauth
from hashlib import sha256
from streamlit.runtime.scriptrunner import add_script_run_ctx
import base64
import requests

# Configure Logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

def render_report_sections(content):
    # 1. Focus on "Actual Problem vs Symptoms"
    problem_pattern = r"(?:Actual Problem vs Symptoms|Problem Statement|Problem vs Symptoms)[\s:]*\n(.*?)(?=\n##|\n#|$)"
    problem_match = re.search(problem_pattern, content, re.DOTALL | re.IGNORECASE)
    
    if problem_match:
        st.markdown("### 🎯 Problem & Symptoms")
        st.info(problem_match.group(1).strip())
        content = content.replace(problem_match.group(0), "")

    # 2. Focus on "Confirmed Root Causes" with scrollable expander
    rc_pattern = r"(?:Confirmed Root Causes|Root Causes)[\s:]*\n(.*?)(?=\n##|\n#|$)"
    rc_match = re.search(rc_pattern, content, re.DOTALL | re.IGNORECASE)
    
    if rc_match:
        with st.expander("🔍 Confirmed Root Causes (Scrollable)", expanded=True):
            st.markdown(f"""
            <div class="scrollable-content">
                {rc_match.group(1).strip()}
            </div>""", unsafe_allow_html=True)
            logger.info("Main Output")
            logger.info(rc_match.group(1).strip())
            # Add a fallback for markdown rendering since HTML div might block markdown parsing
            # if st.button("Render Markdown View"):
            #     st.markdown(rc_match.group(1).strip())
        content = content.replace(rc_match.group(0), "")

    # 3. Remaining sections in standard expanders
    st.markdown("### 📋 Additional Details")
    sections = re.split(r'\n(?=##\s\d+\.\s.+|##\s.+|#\s.+)', content, flags=re.MULTILINE)
    for section in sections:
        section = section.strip()
        if not section or len(section) < 10: continue
            
        lines = section.split('\n')
        first_line = lines[0].strip()
        clean_title = re.sub(r'^#{1,3}\s*', '', first_line)
        body = '\n'.join(lines[1:]).strip()
        
        if body:
            with st.expander(f"**{clean_title}**"):
                # st.markdown(body)
                st.markdown(f"""
                <div class="scrollable-content">
                    {body}
                </div>
                """, unsafe_allow_html=True)
                logger.info(f"Section: {clean_title}\n###################\n")
                logger.info(body)
    # Extract Mermaid diagram if exists
    mermaid_pattern = r"```mermaid\n(.*?)\n```"
    mermaid_match = re.search(mermaid_pattern, content, re.DOTALL)
    
    if mermaid_match:
        mermaid_code = mermaid_match.group(1).strip()
        st.subheader("📊 Download Root Cause Visualizer")
        
        try:
            encoded_mermaid = base64.b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
            image_url = f"https://mermaid.ink/img/{encoded_mermaid}"
            # st.image(image_url, caption="Root Cause Flow Diagram", use_container_width=True)
            
            img_resp = requests.get(image_url)
            if img_resp.status_code == 200:
                st.download_button(
                    label="📥 Download Flow",
                    data=img_resp.content,
                    file_name="rca_flow_diagram.png",
                    mime="image/png",
                    key=f"dl_app_{uuid.uuid4().hex[:8]}"
                )
        except Exception as e:
            st.error(f"Error generating diagram: {str(e)}")
        
        content = re.sub(mermaid_pattern, "", content, flags=re.DOTALL)
        content = re.sub(r'#+\s*\d*\.*\s*.*[Mm]ermaid.*?\n', "", content)


def create_docx(content):
    doc = Document()
    doc.add_heading('Root Cause Analysis Report', 0)
    
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line[0:1].isdigit() and line[1:2] == '.':
            doc.add_paragraph(line[3:], style='List Number')
        else:
            doc.add_paragraph(line)
            
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# Initialize Database
try:
    init_db()
    logger.info("Database initialized successfully.")
except Exception as e:
    logger.error(f"Failed to initialize database: {str(e)}")
    logger.error(traceback.format_exc())
    st.error("Database initialization failed. Check logs for details.")

# --- AUTHENTICATION ---
def hash_pass(password):
    return sha256(password.encode()).hexdigest()

if 'authentication_status' not in st.session_state:
    st.session_state['authentication_status'] = None

# Custom Auth Logic
def login_ui():
    st.title("🔐 Access Control")
    tab_login, tab_signup = st.tabs(["Login", "Sign Up"])
    
    with tab_login:
        with st.form("Login"):
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")
            submit = st.form_submit_button("Sign In")
            if submit:
                user = get_user_by_username(username)
                if user and user.password_hash == hash_pass(password):
                    st.session_state['authentication_status'] = True
                    st.session_state['username'] = username
                    st.session_state['user_id'] = user.id
                    st.rerun()
                else:
                    st.error("Invalid credentials")

    with tab_signup:
        # with st.form("Signup"):
        #     new_user = st.text_input("New Username")
        #     new_email = st.text_input("Email")
        #     new_pass = st.text_input("New Password", type="password")
        #     submit = st.form_submit_button("Register")
        #     if submit:
        #         if get_user_by_username(new_user):
        #             st.error("Username already exists")
        #         else:
        #             create_user(new_user, new_email, hash_pass(new_pass))
        #             st.success("Account created! Please login.")
        st.info("Sign up service is currently unavailable. Please contact the administrator.")
        
if not st.session_state.get('authentication_status'):
    login_ui()
    st.stop()

# --- MAIN APP ---
st.set_page_config(page_title="RCA Intelligence Engine", layout="wide", page_icon="🔍")

# Custom CSS for Premium Look
st.markdown("""
    <style>
    .main { background-color: #0e1117; color: #fafafa; font-family: 'Inter', sans-serif; }
    .stApp { background: radial-gradient(circle at top right, #1a1c24, #0e1117); }
    .stButton>button {
        width: 100%; border-radius: 12px; height: 3.5em;
        background: linear-gradient(135deg, #2e7bcf 0%, #1a5fb4 100%);
        color: white; font-weight: 700; border: none; transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
        text-transform: uppercase; letter-spacing: 1px;
    }
    .stButton>button:hover { transform: translateY(-3px); box-shadow: 0 8px 25px rgba(46, 123, 207, 0.5); }
    
    /* Glassmorphism for containers */
    .stExpander {
        background: rgba(255, 255, 255, 0.03) !important;
        border-radius: 15px !important;
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
        margin-bottom: 1rem !important;
        backdrop-filter: blur(10px);
    }
    .metric-card {
        background: rgba(255, 255, 255, 0.05);
        padding: 20px;
        border-radius: 15px;
        border: 1px solid rgba(255, 255, 255, 0.1);
        text-align: center;
        transition: all 0.3s ease;
    }
    .metric-card:hover { background: rgba(255, 255, 255, 0.08); border-color: #2e7bcf; }
    .status-box {
        background: linear-gradient(90deg, rgba(46, 123, 207, 0.1) 0%, transparent 100%);
        border-left: 5px solid #2e7bcf;
        padding: 20px;
        border-radius: 8px;
    }
    h1, h2, h3 { color: #ffffff !important; font-weight: 800 !important; }
    
    /* Scrollable container for detailed sections */
    .scrollable-content {
        max-height: 300px;
        overflow-y: auto;
        padding-right: 10px;
        scrollbar-width: thin;
        scrollbar-color: #2e7bcf #0e1117;
    }
    .scrollable-content::-webkit-scrollbar { width: 6px; }
    .scrollable-content::-webkit-scrollbar-thumb { background: #2e7bcf; border-radius: 10px; }
    </style>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;600;800&display=swap" rel="stylesheet">
    """, unsafe_allow_html=True)

# Sidebar
with st.sidebar:
    st.title(f"🔍 Hello, {st.session_state['username']}")
    st.markdown("---")
    
    app_mode = st.radio("Navigation", ["Engine", "History Explorer"])
    
    if app_mode == "Engine":
        st.subheader("Sample Scenarios")
        selected_sample = st.selectbox("Select a pre-defined incident:", ["None"] + [s["title"] for s in INCIDENTS])
    
    st.markdown("---")
    if st.button("Logout"):
        st.session_state.clear()
        st.rerun()

if app_mode == "History Explorer":
    st.title("📂 RCA History Explorer")
    search_query = st.text_input("🔍 Search reports by incident description...")
    
    db = SessionLocal()
    query = db.query(FinalReport).filter(FinalReport.user_id == st.session_state['user_id'])
    if search_query:
        query = query.filter(FinalReport.incident_description.ilike(f"%{search_query}%"))
    
    reports = query.order_by(FinalReport.timestamp.desc()).all()
    
    if not reports:
        st.info("No reports found.")
    else:
        for r in reports:
            with st.expander(f"📅 {r.timestamp.strftime('%Y-%m-%d %H:%M')} | {r.incident_description[:100]}..."):
                st.markdown(f"**Tokens**: {r.total_tokens} | **Cost**: ${r.estimated_cost}")
                render_report_sections(r.final_report)
                
    db.close()
    st.stop()

# Main Header
st.title("🚀 RCA Intelligence Engine")
st.caption("Structured Reasoning • Multi-Agent Validation • Actionable Insights")

# Input Section
input_desc = ""
if selected_sample != "None":
    input_desc = next(s["description"] for s in INCIDENTS if s["title"] == selected_sample)

incident_input = st.text_area("Describe the incident in detail...", value=input_desc, height=150)

col_run, col_empty = st.columns([1, 4])
with col_run:
    analyze_btn = st.button("Run Intelligence Engine", type="primary")

# --- BACKGROUND THREADING LOGIC ---
if "workflow_running" not in st.session_state:
    st.session_state.workflow_running = False
if "current_agent" not in st.session_state:
    st.session_state.current_agent = None

def bg_workflow(incident_desc, session_id, user_id):
    st.session_state.workflow_running = True
    try:
        def progress_update(agent_name):
            st.session_state.current_agent = agent_name
            
        messages, tokens, cost = run_rca_workflow(incident_desc, session_id, progress_cb=progress_update)
        
        # Extract final report
        final_msg = ""
        for m in reversed(messages):
            content = m.get("content", "")
            if content and "TERMINATE" in content:
                final_msg = content.replace("TERMINATE", "").strip()
                break
        
        if not final_msg and messages:
            final_msg = next((m.get("content", "") for m in reversed(messages) if m.get("name") != "User_Proxy"), "")

        save_final_report(session_id, incident_desc, final_msg, 
                          user_id=user_id, 
                          tokens=tokens, cost=cost)
        
        st.session_state.final_report = final_msg
        st.session_state.messages = messages
        st.session_state.tokens = tokens
        st.session_state.cost = cost
    except Exception as e:
        st.session_state.workflow_error = str(e)
        logger.error(traceback.format_exc())
    finally:
        st.session_state.workflow_running = False

if analyze_btn and not st.session_state.workflow_running:
    if not incident_input:
        st.error("Please provide an incident description.")
    else:
        session_id = str(uuid.uuid4())
        st.session_state.session_id = session_id
        st.session_state.current_agent = "Incident_Analyst"
        # Start threading
        thread = threading.Thread(target=bg_workflow, args=(incident_input, session_id, st.session_state['user_id']))
        add_script_run_ctx(thread)
        thread.start()

# --- UI POLLING FOR PROGRESS ---
if st.session_state.workflow_running:
    progress_container = st.empty()
    with progress_container.container():
        agent_name = st.session_state.get('current_agent', 'Starting...')
        st.markdown(f'<div class="status-box">Current Phase: <b>{agent_name.replace("_", " ")}</b></div>', unsafe_allow_html=True)
        st.info("The Intelligence Engine is processing in the background. Please wait...")
        # Simple spinner
        st.spinner("Thinking...")
    time.sleep(1)
    st.rerun()

if "workflow_error" in st.session_state:
    st.error(f"Error: {st.session_state.workflow_error}")
    del st.session_state.workflow_error

# Results Section
if "final_report" in st.session_state:
    # --- METRICS BAR ---
    m1, m2, m3 = st.columns(3)
    with m1:
        st.markdown(f'<div class="metric-card"><b>Tokens Used</b><br><span style="font-size:1.5em; color:#2e7bcf;">{st.session_state.get("tokens", 0)}</span></div>', unsafe_allow_html=True)
    with m2:
        st.markdown(f'<div class="metric-card"><b>Est. Cost</b><br><span style="font-size:1.5em; color:#2e7bcf;">${st.session_state.get("cost", 0.0):.4f}</span></div>', unsafe_allow_html=True)
    with m3:
        rating = st.selectbox("Rate Quality (5=Best, 1=Poor):", [5, 4, 3, 2, 1], index=0, key="rating_box", 
                             help="Rate the accuracy and usefulness of the generated RCA.")
        if st.button("Submit Rating"):
            update_report_rating(st.session_state.session_id, rating)
            st.toast("Thank you for your feedback!", icon="✅")

    # st.markdown("---")
    
    # Using Tabs for cleaner UI
    tab_summary, tab_logs = st.tabs(["📋 Final Report", "🕵️ Reasoning Logs"])
    
    with tab_summary:
        with st.container():
            render_report_sections(st.session_state.final_report)
        
        st.markdown("<br>", unsafe_allow_html=True)
        _, col_md, col_docx = st.columns([5, 1, 1])
        # with col_md:
        #     st.download_button(
        #         label="Download Markdown",
        #         data=st.session_state.final_report,
        #         file_name=f"RCA_Report_{st.session_state.session_id}.md",
        #         mime="text/markdown"
        #     )
        with col_docx:
            try:
                docx_data = create_docx(st.session_state.final_report)
                logger.info("DOCX file generated for download.")
                st.download_button(
                    label="Download",
                    data=docx_data,
                    file_name=f"RCA_Report_{st.session_state.session_id}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                logger.error(f"Failed to generate DOCX: {str(e)}")
                st.error("Error generating Word document.")

    with tab_logs:
        st.subheader("Agent Collaboration History")
        for msg in st.session_state.messages:
            role = msg.get("name", "System")
            content = msg.get("content", "")
            if role != "User_Proxy" and content:
                with st.chat_message(role, avatar="🤖" if role != "User_Proxy" else "👤"):
                    st.markdown(f"**{role.replace('_', ' ')}**")
                    st.markdown(content)

# Data Explorer (Optional)
# st.markdown("---")
# if st.checkbox("Show Database Entries"):
#     db = SessionLocal()
#     # ... (code omitted for brevity but remains in your repo if uncommented)
#     db.close()
